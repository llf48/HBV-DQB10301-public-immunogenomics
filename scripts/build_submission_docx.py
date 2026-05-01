from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SUBMISSION = ROOT / "submission"


FILES = [
    (
        SUBMISSION / "HBV_DQB10301_core_gap_submission_manuscript.md",
        SUBMISSION / "HBV_DQB10301_core_gap_submission_manuscript.docx",
        "manuscript",
    ),
    (
        SUBMISSION / "HBV_DQB10301_cover_letter.md",
        SUBMISSION / "HBV_DQB10301_cover_letter.docx",
        "letter",
    ),
    (
        SUBMISSION / "HBV_DQB10301_figure_legends.md",
        SUBMISSION / "HBV_DQB10301_figure_legends.docx",
        "legends",
    ),
    (
        SUBMISSION / "HBV_DQB10301_supplementary_note.md",
        SUBMISSION / "HBV_DQB10301_supplementary_note.docx",
        "supplement",
    ),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(clean_inline(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)


def clean_inline(text: str) -> str:
    text = text.replace("**", "")
    text = text.replace("`", "")
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
    return text


def add_inline_runs(paragraph, text: str) -> None:
    """Add a paragraph with light Markdown emphasis handling."""
    pos = 0
    for match in re.finditer(r"\*\*(.+?)\*\*|`(.+?)`", text):
        if match.start() > pos:
            run = paragraph.add_run(clean_inline(text[pos : match.start()]))
            run.font.name = "Times New Roman"
        if match.group(1):
            run = paragraph.add_run(clean_inline(match.group(1)))
            run.bold = True
            run.font.name = "Times New Roman"
        else:
            run = paragraph.add_run(clean_inline(match.group(2)))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(clean_inline(text[pos:]))
        run.font.name = "Times New Roman"


def style_document(doc: Document, kind: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    if kind == "letter":
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.6)
        normal.font.size = Pt(10.5)
        normal.paragraph_format.space_after = Pt(3)
        normal.paragraph_format.line_spacing = 1.0

    for name, size, color in [
        ("Title", 18, "1F4E79"),
        ("Heading 1", 14, "1F4E79"),
        ("Heading 2", 12, "3F6C51"),
        ("Heading 3", 11, "3F6C51"),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    if kind == "manuscript":
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("HBV DQB1*03:01 core presentation gap")
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)


def is_table_block(lines: list[str], start: int) -> bool:
    return (
        start + 1 < len(lines)
        and lines[start].lstrip().startswith("|")
        and re.match(r"^\s*\|?[-:\s|]+\|?\s*$", lines[start + 1]) is not None
    )


def add_table(doc: Document, table_lines: list[str]) -> None:
    rows = []
    for line in table_lines:
        stripped = line.strip()
        if re.match(r"^\|?[-:\s|]+\|?$", stripped):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return

    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for row_idx, row in enumerate(rows):
        for col_idx in range(max_cols):
            cell = table.cell(row_idx, col_idx)
            text = row[col_idx] if col_idx < len(row) else ""
            set_cell_text(cell, text, bold=row_idx == 0)
            if row_idx == 0:
                set_cell_shading(cell, "D9EAF7")

    doc.add_paragraph()


def add_markdown_file(doc: Document, md_path: Path, kind: str) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        raw = lines[i].rstrip()

        if raw.startswith("```"):
            if in_code:
                para = doc.add_paragraph()
                para.style = doc.styles["Normal"]
                run = para.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(8)
                code_lines = []
                in_code = False
            else:
                in_code = True
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        if not raw.strip():
            i += 1
            continue

        if is_table_block(lines, i):
            block: list[str] = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                block.append(lines[i])
                i += 1
            add_table(doc, block)
            continue

        if raw.startswith("# "):
            para = doc.add_paragraph()
            para.style = doc.styles["Title"]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if kind != "supplement" else WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(para, raw[2:])
        elif raw.startswith("## "):
            heading = clean_inline(raw[3:])
            if kind == "manuscript" and heading == "References":
                doc.add_page_break()
            doc.add_heading(heading, level=1)
        elif raw.startswith("### "):
            doc.add_heading(clean_inline(raw[4:]), level=2)
        elif raw.startswith("#### "):
            doc.add_heading(clean_inline(raw[5:]), level=3)
        elif re.match(r"^\d+\.\s+", raw):
            para = doc.add_paragraph(style="List Number")
            add_inline_runs(para, re.sub(r"^\d+\.\s+", "", raw))
        elif raw.startswith("- "):
            para = doc.add_paragraph(style="List Bullet")
            add_inline_runs(para, raw[2:])
        elif raw.startswith("> "):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.right_indent = Inches(0.25)
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(8)
            add_inline_runs(para, raw[2:])
        else:
            para = doc.add_paragraph()
            add_inline_runs(para, raw)

        i += 1


def build_docx(md_path: Path, out_path: Path, kind: str) -> None:
    doc = Document()
    style_document(doc, kind)
    add_markdown_file(doc, md_path, kind)
    doc.save(out_path)


def main() -> None:
    for md_path, out_path, kind in FILES:
        build_docx(md_path, out_path, kind)
        print(out_path)


if __name__ == "__main__":
    main()
